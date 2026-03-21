"""
app/parsers/resume_parser.py
─────────────────────────────
Core resume parsing engine.

Strategy:
  PDF  → Try PyMuPDF first (fastest, preserves layout)
         → Fall back to pdfminer.six (better for scanned-text hybrids)
  DOCX → python-docx (paragraph + table extraction)

After text extraction, a lightweight NLP pipeline runs to detect
sections and hand off to entity extractors.
"""

import io
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ── Optional heavy imports (graceful degradation) ────────
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF not installed. PDF fallback to pdfminer.")

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    HAS_PDFMINER = True
except ImportError:
    HAS_PDFMINER = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ─────────────────────────────────────────────────────────
# Section header patterns (regex)
# ─────────────────────────────────────────────────────────
SECTION_PATTERNS: Dict[str, List[str]] = {
    "summary":        [r"\b(summary|objective|profile|about me|professional summary)\b"],
    "skills":         [r"\b(skills?|technical skills?|core competencies|expertise|technologies)\b"],
    "experience":     [r"\b(experience|work experience|employment|professional experience|career)\b"],
    "education":      [r"\b(education|academic|qualification|degree|university|college)\b"],
    "projects":       [r"\b(projects?|personal projects?|academic projects?|portfolio)\b"],
    "certifications": [r"\b(certifications?|certificates?|licenses?|accreditations?)\b"],
    "languages":      [r"\b(languages?|language proficiency|spoken languages?)\b"],
    "achievements":   [r"\b(achievements?|awards?|honors?|accomplishments?)\b"],
}


class ResumeParser:
    """
    Orchestrates resume text extraction and structured data parsing.

    Usage:
        parser = ResumeParser()
        data = parser.parse(file_bytes=b"...", file_type="pdf")
    """

    def __init__(self):
        # Lazy-load spaCy — only needed for NER, loaded once
        self._nlp = None

    @property
    def nlp(self):
        """Lazy load spaCy model (en_core_web_lg or sm as fallback)."""
        if self._nlp is None:
            import spacy
            from app.core.config import settings
            try:
                self._nlp = spacy.load(settings.SPACY_MODEL)
            except OSError:
                logger.warning(f"{settings.SPACY_MODEL} not found. Falling back to en_core_web_sm")
                self._nlp = spacy.load("en_core_web_sm")
        return self._nlp

    # ──────────────────────────────────────────────────────
    # Public entry point
    # ──────────────────────────────────────────────────────

    def parse(self, file_bytes: bytes, file_type: str) -> Dict:
        """
        Main method: extract text → detect sections → extract entities.

        Returns a dict matching ResumeDocument fields.
        Raises ResumeParseError / EmptyResumeError on failure.
        """
        from app.core.exceptions import EmptyResumeError, ResumeParseError

        # Step 1 – Extract raw text
        file_type = file_type.lower().strip(".")
        if file_type == "pdf":
            raw_text = self._extract_pdf(file_bytes)
        elif file_type in ("docx", "doc"):
            raw_text = self._extract_docx(file_bytes)
        else:
            raise ResumeParseError(f"Unsupported file type: {file_type}")

        if not raw_text or len(raw_text.strip()) < 50:
            raise EmptyResumeError()

        # Step 2 – Detect sections (returns {section_name: text_block})
        sections = self._detect_sections(raw_text)

        # Step 3 – Extract each field
        parsed = {
            "raw_text": raw_text,
            "word_count": len(raw_text.split()),
            "char_count": len(raw_text),
            "section_flags": {k: bool(v) for k, v in sections.items()},
            **self._extract_contact_info(raw_text),
            "summary":        self._extract_summary(sections),
            "skills":         self._extract_skills(sections, raw_text),
            "soft_skills":    self._extract_soft_skills(raw_text),
            "experience":     self._extract_experience(sections),
            "education":      self._extract_education(sections),
            "projects":       self._extract_projects(sections),
            "certifications": self._extract_certifications(sections),
            "languages":      self._extract_languages(sections),
        }

        logger.info(f"Parsed resume: {parsed.get('name','?')} | {parsed['word_count']} words | sections: {list(sections.keys())}")
        return parsed

    # ──────────────────────────────────────────────────────
    # Text extraction
    # ──────────────────────────────────────────────────────

    def _extract_pdf(self, file_bytes: bytes) -> str:
        """PyMuPDF preferred → pdfminer fallback."""
        text = ""

        if HAS_PYMUPDF:
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                pages = [page.get_text("text") for page in doc]
                text = "\n".join(pages)
                doc.close()
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {e}. Trying pdfminer.")
                text = ""

        if not text.strip() and HAS_PDFMINER:
            try:
                text = pdfminer_extract(io.BytesIO(file_bytes))
            except Exception as e:
                logger.error(f"pdfminer failed: {e}")

        return self._clean_text(text)

    def _extract_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX preserving paragraph order."""
        if not HAS_DOCX:
            from app.core.exceptions import ResumeParseError
            raise ResumeParseError("python-docx not installed")

        doc = DocxDocument(io.BytesIO(file_bytes))
        lines = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())

        # Extract table cells (skills often in tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    lines.append(row_text)

        return self._clean_text("\n".join(lines))

    # ──────────────────────────────────────────────────────
    # Section detection
    # ──────────────────────────────────────────────────────

    def _detect_sections(self, text: str) -> Dict[str, str]:
        """
        Split resume text into named sections.

        Finds header lines that match SECTION_PATTERNS and collects
        the text until the next header.

        Returns {section_name: section_text}
        """
        lines = text.split("\n")
        sections: Dict[str, str] = {}
        current_section = "header"
        current_lines: List[str] = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                current_lines.append("")
                continue

            matched_section = self._match_section_header(stripped)

            if matched_section and len(stripped) < 60:
                # Save previous section
                if current_lines:
                    sections[current_section] = "\n".join(current_lines).strip()
                current_section = matched_section
                current_lines = []
            else:
                current_lines.append(stripped)

        # Save last section
        if current_lines:
            sections[current_section] = "\n".join(current_lines).strip()

        return sections

    def _match_section_header(self, line: str) -> Optional[str]:
        """Return section name if line matches a known header pattern."""
        line_lower = line.lower()
        for section, patterns in SECTION_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, line_lower):
                    return section
        return None

    # ──────────────────────────────────────────────────────
    # Entity extraction
    # ──────────────────────────────────────────────────────

    def _extract_contact_info(self, text: str) -> Dict:
        """Extract name, email, phone, LinkedIn, GitHub from top ~30 lines."""
        header_text = "\n".join(text.split("\n")[:30])

        return {
            "name":     self._extract_name(header_text),
            "email":    self._extract_email(header_text),
            "phone":    self._extract_phone(header_text),
            "linkedin": self._extract_linkedin(header_text),
            "github":   self._extract_github(header_text),
            "location": self._extract_location(header_text),
        }

    def _extract_name(self, text: str) -> str:
        """
        Name is usually the first non-empty line of a resume.
        Falls back to spaCy PERSON entity if first line looks wrong.
        """
        first_lines = [l.strip() for l in text.split("\n") if l.strip()]
        if first_lines:
            candidate = first_lines[0]
            # Likely a name: 2-4 words, no @, no digits, not too long
            if (
                2 <= len(candidate.split()) <= 5
                and "@" not in candidate
                and not re.search(r"\d", candidate)
                and len(candidate) < 50
            ):
                return candidate

        # spaCy fallback
        doc = self.nlp(text[:500])
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                return ent.text
        return ""

    def _extract_email(self, text: str) -> str:
        emails = re.findall(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}", text)
        return emails[0] if emails else ""

    def _extract_phone(self, text: str) -> str:
        # Matches common formats: +91-9876543210, (123) 456-7890, 1234567890
        phones = re.findall(
            r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text
        )
        return phones[0] if phones else ""

    def _extract_linkedin(self, text: str) -> Optional[str]:
        match = re.search(r"linkedin\.com/in/([\w-]+)", text, re.IGNORECASE)
        return f"linkedin.com/in/{match.group(1)}" if match else None

    def _extract_github(self, text: str) -> Optional[str]:
        match = re.search(r"github\.com/([\w-]+)", text, re.IGNORECASE)
        return f"github.com/{match.group(1)}" if match else None

    def _extract_location(self, text: str) -> Optional[str]:
        # Look for City, State/Country patterns in first 10 lines
        lines = text.split("\n")[:10]
        location_pattern = re.compile(
            r"\b([A-Z][a-zA-Z\s]+),\s*([A-Z]{2}|[A-Z][a-zA-Z\s]+)\b"
        )
        for line in lines:
            m = location_pattern.search(line)
            if m:
                return m.group(0)
        return None

    def _extract_summary(self, sections: Dict[str, str]) -> Optional[str]:
        return sections.get("summary") or sections.get("header", "")[:500] or None

    def _extract_skills(self, sections: Dict[str, str], raw_text: str) -> List[str]:
        """
        Extract skills from the skills section.
        Also scans full text for known tech keywords as backup.
        """
        skills_text = sections.get("skills", "")
        if not skills_text:
            # No dedicated section — scan full text
            skills_text = raw_text

        # Split on common delimiters: comma, pipe, bullet, newline
        raw_skills = re.split(r"[,|•\n·]+", skills_text)
        skills = []
        for s in raw_skills:
            s = s.strip().strip("•-–—*").strip()
            # Keep tokens 2–40 chars that are "skill-like"
            if 2 <= len(s) <= 40 and not re.match(r"^\d+$", s):
                skills.append(s)

        # Deduplicate preserving order (case-insensitive)
        seen = set()
        unique = []
        for skill in skills:
            k = skill.lower()
            if k not in seen:
                seen.add(k)
                unique.append(skill)

        return unique[:80]  # cap at 80 skills

    def _extract_soft_skills(self, text: str) -> List[str]:
        SOFT_SKILLS = [
            "leadership", "communication", "teamwork", "problem solving",
            "critical thinking", "adaptability", "creativity", "time management",
            "collaboration", "presentation", "analytical", "attention to detail",
            "project management", "negotiation", "mentoring",
        ]
        found = []
        text_lower = text.lower()
        for skill in SOFT_SKILLS:
            if skill in text_lower:
                found.append(skill.title())
        return found

    def _extract_experience(self, sections: Dict[str, str]) -> List[Dict]:
        """
        Parse work experience blocks.
        Each block is: Company | Role | Date range | Bullets
        """
        exp_text = sections.get("experience", "")
        if not exp_text:
            return []

        experiences = []
        # Split on blank lines or date patterns as block separators
        blocks = re.split(r"\n{2,}", exp_text)

        for block in blocks:
            if not block.strip():
                continue
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue

            entry = {
                "company": "",
                "role": "",
                "start_date": None,
                "end_date": None,
                "description": "",
                "technologies": [],
            }

            # Date pattern: Jan 2022 – Dec 2023, 2020-2022, Present
            date_pattern = re.compile(
                r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|[12]\d{3})"
                r"[\s,–-]+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|"
                r"[12]\d{3}|Present|Current|Now)",
                re.IGNORECASE,
            )

            desc_lines = []
            for i, line in enumerate(lines):
                date_match = date_pattern.search(line)
                if date_match and i < 4:
                    entry["start_date"] = date_match.group(1)
                    entry["end_date"] = date_match.group(2)
                elif i == 0:
                    # First line is usually role or company
                    entry["role"] = line
                elif i == 1 and not date_match:
                    entry["company"] = line
                else:
                    desc_lines.append(line)

            entry["description"] = "\n".join(desc_lines)

            # Extract mentioned technologies from description
            entry["technologies"] = self._extract_tech_from_text(entry["description"])

            if entry["role"] or entry["company"]:
                experiences.append(entry)

        return experiences

    def _extract_education(self, sections: Dict[str, str]) -> List[Dict]:
        edu_text = sections.get("education", "")
        if not edu_text:
            return []

        educations = []
        blocks = re.split(r"\n{2,}", edu_text)
        degree_pattern = re.compile(
            r"\b(B\.?Tech|M\.?Tech|B\.?E|B\.?Sc|M\.?Sc|MBA|PhD|MCA|BCA|"
            r"Bachelor|Master|Doctorate|Associate|Diploma)\b",
            re.IGNORECASE,
        )
        year_pattern = re.compile(r"\b(19|20)\d{2}\b")

        for block in blocks:
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue

            entry = {"institution": "", "degree": "", "field": "", "year": None, "gpa": None}
            for line in lines:
                deg_match = degree_pattern.search(line)
                year_match = year_pattern.search(line)
                gpa_match = re.search(r"(GPA|CGPA|CPI)[:\s]*([\d.]+)", line, re.IGNORECASE)

                if deg_match:
                    entry["degree"] = deg_match.group(0)
                    # Field of study is usually after the degree keyword
                    rest = line[deg_match.end():].strip("., :")
                    if rest and len(rest) < 60:
                        entry["field"] = rest
                elif year_match and not entry["year"]:
                    entry["year"] = year_match.group(0)
                elif gpa_match:
                    entry["gpa"] = gpa_match.group(2)
                elif not entry["institution"] and len(line) > 5:
                    entry["institution"] = line

            if entry["institution"] or entry["degree"]:
                educations.append(entry)

        return educations

    def _extract_projects(self, sections: Dict[str, str]) -> List[Dict]:
        proj_text = sections.get("projects", "")
        if not proj_text:
            return []

        projects = []
        blocks = re.split(r"\n{2,}", proj_text)
        url_pattern = re.compile(r"https?://[^\s]+|github\.com/[^\s]+", re.IGNORECASE)

        for block in blocks:
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue

            name = lines[0].strip("•-–—*:").strip()
            description = " ".join(lines[1:])
            url_match = url_pattern.search(block)

            projects.append({
                "name": name,
                "description": description,
                "technologies": self._extract_tech_from_text(description),
                "url": url_match.group(0) if url_match else None,
            })

        return projects

    def _extract_certifications(self, sections: Dict[str, str]) -> List[str]:
        cert_text = sections.get("certifications", "")
        if not cert_text:
            return []
        lines = [l.strip("•-–—*:").strip() for l in cert_text.split("\n") if l.strip()]
        return [l for l in lines if len(l) > 3]

    def _extract_languages(self, sections: Dict[str, str]) -> List[str]:
        lang_text = sections.get("languages", "")
        if not lang_text:
            return []
        return [l.strip() for l in re.split(r"[,|•\n·]+", lang_text) if l.strip()]

    # ──────────────────────────────────────────────────────
    # Helpers
    # ──────────────────────────────────────────────────────

    def _extract_tech_from_text(self, text: str) -> List[str]:
        """Match known tech keywords from free text."""
        TECH_KEYWORDS = {
            "python", "java", "javascript", "typescript", "c++", "c#", "go",
            "rust", "react", "vue", "angular", "node.js", "django", "fastapi",
            "flask", "spring", "kubernetes", "docker", "aws", "gcp", "azure",
            "postgresql", "mysql", "mongodb", "redis", "kafka", "rabbitmq",
            "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy",
            "git", "linux", "terraform", "graphql", "rest", "grpc",
        }
        text_lower = text.lower()
        return [kw for kw in TECH_KEYWORDS if kw in text_lower]

    def _clean_text(self, text: str) -> str:
        """Normalise whitespace and remove non-printable characters."""
        # Remove non-printable (keep newlines)
        text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)
        # Collapse 3+ consecutive newlines to 2
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Collapse multiple spaces
        text = re.sub(r"[ \t]{2,}", " ", text)
        return text.strip()
